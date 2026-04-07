import re
import tempfile
from fpdf import FPDF


def sanitize_text(text: str) -> str:
    """Replace Unicode characters that Courier can't render with ASCII equivalents."""
    replacements = {
        "\u2014": "--",   # em dash
        "\u2013": "-",    # en dash
        "\u2018": "'",    # left single quote
        "\u2019": "'",    # right single quote
        "\u201c": '"',    # left double quote
        "\u201d": '"',    # right double quote
        "\u2026": "...",  # ellipsis
        "\u2022": "-",    # bullet
        "\u00a0": " ",    # non-breaking space
        "\u2032": "'",    # prime
        "\u2033": '"',    # double prime
        "\u00b7": "-",    # middle dot
        "\u200b": "",     # zero-width space
        "\u00e9": "e",    # e-acute
        "\u00e8": "e",    # e-grave
        "\u00e0": "a",    # a-grave
        "\u2192": "->",   # right arrow
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    # Strip any remaining non-latin1 characters
    text = text.encode("latin-1", errors="replace").decode("latin-1")
    return text


# ═══════════════════════════════════════════════════
# SCREENPLAY PDF (matches FINAL_ reference format)
# ═══════════════════════════════════════════════════

class ScreenplayPDF(FPDF):
    """PDF formatted like a proper screenplay with running header."""

    def __init__(self, script_title=""):
        super().__init__()
        self._script_title = script_title
        self._on_title_page = True
        self.add_page()
        self.set_auto_page_break(auto=True, margin=25)
        self.set_margins(25, 25, 25)

    def header(self):
        if self._on_title_page:
            return
        # Running header: bold underlined title
        self.set_font("Courier", "BU", 10)
        title_text = sanitize_text(self._script_title.upper()) if self._script_title else ""
        self.cell(0, 10, title_text, new_x="LMARGIN", new_y="NEXT")
        self.ln(4)

    def footer(self):
        if self._on_title_page:
            return
        self.set_y(-20)
        self.set_font("Courier", "", 10)
        self.cell(0, 10, f"{self.page_no() - 1}.", align="R")

    def title_page(self, title, logline=""):
        self._on_title_page = True
        self.ln(80)
        self.set_font("Courier", "B", 24)
        self.cell(0, 15, sanitize_text(title.upper()), align="C", new_x="LMARGIN", new_y="NEXT")
        if logline:
            self.ln(10)
            self.set_font("Courier", "", 12)
            self.multi_cell(0, 6, sanitize_text(logline), align="C")
        self._on_title_page = False
        self.add_page()

    def scene_heading(self, text):
        self.ln(4)
        self.set_font("Courier", "BU", 12)
        self.cell(0, 6, sanitize_text(text.upper()), new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def action_line(self, text):
        self.set_font("Courier", "", 12)
        self.multi_cell(0, 6, sanitize_text(text))
        self.ln(2)

    def character_name(self, name):
        self.ln(2)
        self.set_font("Courier", "B", 12)
        self.cell(0, 6, sanitize_text(name.upper()), align="C", new_x="LMARGIN", new_y="NEXT")

    def parenthetical(self, text):
        self.set_font("Courier", "I", 12)
        x_offset = 55
        self.set_x(x_offset)
        self.cell(0, 6, sanitize_text(f"({text})"), new_x="LMARGIN", new_y="NEXT")

    def dialogue(self, text):
        self.set_font("Courier", "", 12)
        left_margin = 40
        right_margin = 40
        self.set_left_margin(left_margin)
        self.set_right_margin(right_margin)
        self.set_x(left_margin)
        self.multi_cell(0, 6, sanitize_text(text))
        self.set_left_margin(25)
        self.set_right_margin(25)
        self.ln(1)

    def clip_marker(self, text):
        self.set_font("Courier", "B", 10)
        self.set_text_color(180, 40, 40)
        self.cell(0, 6, sanitize_text(text), new_x="LMARGIN", new_y="NEXT")
        self.set_text_color(0, 0, 0)

    def series_heading(self, text):
        """For BEGIN SERIES OF SHOTS / END SERIES OF SHOTS and sub-headings."""
        self.ln(4)
        self.set_font("Courier", "BU", 12)
        self.cell(0, 6, sanitize_text(text.upper()), new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def transition(self, text):
        """For CUT TO BLACK, FADE OUT, etc."""
        self.ln(4)
        self.set_font("Courier", "", 12)
        self.cell(0, 6, sanitize_text(text.upper()), align="R", new_x="LMARGIN", new_y="NEXT")
        self.ln(2)


def _extract_section(script_text, section_name, next_sections=None):
    """Extract a section from the script text by header name."""
    if next_sections is None:
        next_sections = []
    # Build pattern for end boundary
    end_patterns = "|".join([re.escape(f"## {s}") for s in next_sections]) if next_sections else "$"
    pattern = rf"##\s*{re.escape(section_name)}\s*\n(.*?)(?={end_patterns}|$)"
    match = re.search(pattern, script_text, re.DOTALL | re.IGNORECASE)
    return match.group(1).strip() if match else ""


def parse_script_to_pdf(script_text: str) -> bytes:
    """Parse a screenplay-formatted script and generate a PDF."""
    # Extract title and logline
    title = "Untitled Script"
    logline = ""

    title_match = re.search(r"(?:TITLE|Title)[:\s]*\n*(.+?)(?:\n|$)", script_text)
    if title_match:
        title = title_match.group(1).strip().strip("#").strip()

    logline_match = re.search(
        r"(?:LOGLINE|Logline)[:\s]*\n*(.+?)(?:\n\n|\n#|$)", script_text, re.DOTALL
    )
    if logline_match:
        logline = logline_match.group(1).strip()

    # Build header text like "FINAL:TITLE aka SUBTITLE"
    header_title = f"FINAL:{title}"

    pdf = ScreenplayPDF(script_title=header_title)
    pdf.title_page(title, logline)

    # Find the script section
    script_section = script_text
    script_match = re.search(
        r"##\s*SCRIPT\s*\n(.*?)(?=##\s*CLIP|$)", script_text, re.DOTALL
    )
    if script_match:
        script_section = script_match.group(1)

    lines = script_section.strip().split("\n")

    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue

        # Clip markers
        if "[CLIP START]" in stripped or "[CLIP END]" in stripped:
            pdf.clip_marker(stripped)
            continue

        # Series of shots headings
        if re.match(r"^(BEGIN\s+SERIES|END\s+SERIES|SERIES\s+OF)", stripped, re.IGNORECASE):
            pdf.series_heading(stripped)
            continue

        # Scene headings: INT. / EXT.
        if re.match(r"^(INT\.|EXT\.|INT/EXT\.)", stripped, re.IGNORECASE):
            pdf.scene_heading(stripped)
            continue

        # Sub-headings within series of shots (e.g., "FRONT DOOR - DAY 3:")
        if re.match(r"^[A-Z][A-Z\s/]+-\s*(DAY|NIGHT|LATER|CONTINUOUS)", stripped) and stripped.endswith(":"):
            pdf.series_heading(stripped.rstrip(":"))
            continue

        # Transitions (CUT TO, FADE, etc.)
        if re.match(r"^(CUT TO|FADE TO|FADE OUT|FADE IN|SMASH CUT|CUT TO BLACK)", stripped, re.IGNORECASE):
            pdf.transition(stripped)
            continue

        # Character names: all caps line, possibly with (V.O.) or (O.S.) or (CONT'D)
        if re.match(r"^[A-Z][A-Z\s\-\.]+(\s*\(.*?\))?\s*$", stripped) and len(stripped) < 40:
            pdf.character_name(stripped)
            continue

        # Parentheticals: lines in parentheses
        if re.match(r"^\(.*\)$", stripped):
            pdf.parenthetical(stripped.strip("()"))
            continue

        # INSERT lines
        if stripped.startswith("INSERT"):
            pdf.action_line(stripped)
            continue

        # Stage directions in brackets
        if re.match(r"^\[.*\]$", stripped):
            pdf.action_line(stripped.strip("[]"))
            continue

        # (MORE) continuation
        if stripped == "(MORE)":
            pdf.set_font("Courier", "", 10)
            pdf.cell(0, 6, sanitize_text("(MORE)"), align="C", new_x="LMARGIN", new_y="NEXT")
            continue

        # Default: treat as action/description
        pdf.action_line(stripped)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(tmp.name)
    tmp.seek(0)
    with open(tmp.name, "rb") as f:
        return f.read()


# ═══════════════════════════════════════════════════
# CASTING BREAKDOWN PDF (matches CB_ reference format)
# ═══════════════════════════════════════════════════

class CastingBreakdownPDF(FPDF):
    """PDF formatted like a professional casting breakdown."""

    def __init__(self):
        super().__init__()
        self.add_page()
        self.set_auto_page_break(auto=True, margin=20)
        self.set_margins(20, 20, 20)

    def section_divider(self):
        self.ln(2)
        self.set_draw_color(0, 0, 0)
        self.set_line_width(0.5)
        self.line(20, self.get_y(), self.w - 20, self.get_y())
        self.ln(4)

    def bold_heading(self, text):
        self.set_font("Helvetica", "BU", 12)
        self.cell(0, 8, sanitize_text(text), new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def bold_label_value(self, label, value):
        self.set_font("Helvetica", "B", 11)
        label_text = sanitize_text(label)
        label_w = self.get_string_width(label_text) + 2
        self.cell(label_w, 6, label_text)
        self.set_font("Helvetica", "", 11)
        # Use cell for short values to avoid width issues
        remaining = self.w - self.get_x() - self.r_margin
        value_text = sanitize_text(value)
        if self.get_string_width(value_text) < remaining:
            self.cell(0, 6, value_text, new_x="LMARGIN", new_y="NEXT")
        else:
            self.ln()
            self.multi_cell(0, 6, value_text)

    def body_text(self, text):
        self.set_font("Helvetica", "", 11)
        self.multi_cell(0, 6, sanitize_text(text))
        self.ln(2)

    def bullet_item(self, text):
        self.set_font("Helvetica", "", 11)
        indent = self.l_margin + 6
        self.set_x(indent)
        self.cell(6, 6, sanitize_text("-"))
        self.set_x(indent + 6)
        self.multi_cell(self.w - indent - 6 - self.r_margin, 6, sanitize_text(text))

    def sub_bullet_item(self, text):
        self.set_font("Helvetica", "", 10)
        indent = self.l_margin + 16
        self.set_x(indent)
        self.cell(6, 6, sanitize_text("o"))
        self.set_x(indent + 6)
        self.multi_cell(self.w - indent - 6 - self.r_margin, 6, sanitize_text(text))


def parse_casting_to_pdf(script_text: str) -> bytes:
    """Parse the casting breakdown section and generate a formatted PDF."""
    pdf = CastingBreakdownPDF()

    # Extract title
    title = "Untitled Script"
    title_match = re.search(r"(?:TITLE|Title)[:\s]*\n*(.+?)(?:\n|$)", script_text)
    if title_match:
        title = title_match.group(1).strip().strip("#").strip()

    # Extract logline/synopsis
    logline = ""
    logline_match = re.search(
        r"(?:LOGLINE|Logline)[:\s]*\n*(.+?)(?:\n\n|\n#|$)", script_text, re.DOTALL
    )
    if logline_match:
        logline = logline_match.group(1).strip()

    # Project title
    pdf.set_font("Helvetica", "BU", 14)
    pdf.cell(0, 10, sanitize_text(f"NAME OF PROJECT: {title.upper()}"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Synopsis
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, sanitize_text("SYNOPSIS/STORYLINE:"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    if logline:
        pdf.body_text(logline)

    pdf.section_divider()

    # Details section (placeholder — user can fill in)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, sanitize_text("DETAILS:"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.bold_label_value("SHOOTING DATE: ", "TBD")
    pdf.bold_label_value("LOCATION: ", "TBD")
    pdf.bold_label_value("VIRTUAL ZOOM TABLE READ: ", "TBD")
    pdf.ln(2)
    for detail in [
        "Pay: TBD",
        "Snacks and Drinks Provided",
        "Lunch not provided",
        "Demo Reel Footage Provided (upon request)",
        "Actors must arrange their own transportation",
        "Actors should arrive camera-ready with wardrobe, hair, and makeup complete",
    ]:
        pdf.bullet_item(detail)
    pdf.ln(2)

    pdf.section_divider()

    # Character descriptions
    pdf.bold_heading("CHARACTER DESCRIPTIONS:")
    pdf.ln(2)

    casting_section = _extract_section(
        script_text, "CASTING BREAKDOWN",
        ["PROPS LIST", "LOCATION LIST"]
    )

    if casting_section:
        lines = casting_section.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line or line.startswith("---"):
                i += 1
                continue

            # Character heading: **NAME (Role)** or bold character name
            char_match = re.match(r"^\*\*(.+?)\*\*\s*$", line)
            if char_match:
                char_name = char_match.group(1).strip()
                pdf.section_divider()
                pdf.set_font("Helvetica", "B", 12)
                pdf.cell(0, 8, sanitize_text(char_name), new_x="LMARGIN", new_y="NEXT")
                pdf.ln(2)
                i += 1
                continue

            # Bold label: **Key:** value
            label_match = re.match(r"^\*\*(.+?):\*\*\s*(.*)", line)
            if label_match:
                label = label_match.group(1).strip()
                value = label_match.group(2).strip()
                # Multi-line value: gather continuation lines
                while i + 1 < len(lines) and lines[i + 1].strip() and not lines[i + 1].strip().startswith("**") and not lines[i + 1].strip().startswith("- ") and not lines[i + 1].strip().startswith("---"):
                    i += 1
                    value += " " + lines[i].strip()
                pdf.bold_label_value(f"{label}: ", value)
                pdf.ln(1)
                i += 1
                continue

            # Bullet items (day-by-day breakdowns)
            if line.startswith("- "):
                content = line[2:].strip()
                # Remove markdown bold markers
                content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
                pdf.bullet_item(content)
                i += 1
                continue

            # Sub-bullets
            if re.match(r"^\s+-\s", line) or line.startswith("  -") or line.startswith("    -"):
                content = re.sub(r"^\s*-\s*", "", line).strip()
                content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
                pdf.sub_bullet_item(content)
                i += 1
                continue

            # Default text line
            if line:
                pdf.body_text(re.sub(r"\*\*(.+?)\*\*", r"\1", line))

            i += 1

    # Props List
    props_section = _extract_section(script_text, "PROPS LIST", ["LOCATION LIST", "CASTING BREAKDOWN"])
    if props_section:
        pdf.section_divider()
        pdf.bold_heading("PROPS LIST:")
        pdf.ln(2)
        for line in props_section.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("- "):
                content = line[2:].strip()
                content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
                pdf.bullet_item(content)
            elif line.startswith("* "):
                content = line[2:].strip()
                content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
                pdf.bullet_item(content)
            else:
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                pdf.body_text(clean)

    # Location List
    location_section = _extract_section(script_text, "LOCATION LIST", ["CASTING BREAKDOWN", "PROPS LIST"])
    if location_section:
        pdf.section_divider()
        pdf.bold_heading("LOCATION LIST:")
        pdf.ln(2)
        for line in location_section.split("\n"):
            line = line.strip()
            if not line:
                continue
            # Scene heading lines
            scene_match = re.match(r"^\*\*(.+?)\*\*\s*$", line)
            if scene_match:
                pdf.ln(2)
                pdf.set_font("Helvetica", "B", 11)
                pdf.cell(0, 7, sanitize_text(scene_match.group(1)), new_x="LMARGIN", new_y="NEXT")
                pdf.ln(1)
                continue
            if line.startswith("- "):
                content = line[2:].strip()
                content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
                pdf.bullet_item(content)
            else:
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                pdf.body_text(clean)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(tmp.name)
    tmp.seek(0)
    with open(tmp.name, "rb") as f:
        return f.read()


# ═══════════════════════════════════════════════════
# CLIP BREAKDOWN PDF + DOCX
# ═══════════════════════════════════════════════════

def parse_clip_breakdown_to_pdf(script_text: str) -> bytes:
    """Extract clip breakdown section and generate a PDF."""
    title = "Untitled Script"
    title_match = re.search(r"(?:TITLE|Title)[:\s]*\n*(.+?)(?:\n|$)", script_text)
    if title_match:
        title = title_match.group(1).strip().strip("#").strip()

    clip_section = _extract_section(
        script_text, "CLIP BREAKDOWN",
        ["CASTING BREAKDOWN", "PROPS LIST", "LOCATION LIST"]
    )

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_margins(20, 20, 20)

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, sanitize_text(f"CLIP BREAKDOWN"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 12)
    pdf.cell(0, 8, sanitize_text(title), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(8)

    if clip_section:
        for line in clip_section.split("\n"):
            line = line.strip()
            if not line:
                pdf.ln(3)
                continue
            # Bold items (clip titles)
            bold_match = re.match(r"^\*\*(.+?)\*\*:?\s*(.*)", line)
            if bold_match:
                pdf.set_font("Helvetica", "B", 11)
                pdf.cell(0, 7, sanitize_text(bold_match.group(1)), new_x="LMARGIN", new_y="NEXT")
                if bold_match.group(2):
                    pdf.set_font("Helvetica", "", 11)
                    pdf.multi_cell(0, 6, sanitize_text(bold_match.group(2)))
                pdf.ln(2)
                continue
            # Numbered items
            num_match = re.match(r"^(\d+)\.\s*(.*)", line)
            if num_match:
                pdf.set_font("Helvetica", "B", 11)
                indent = pdf.l_margin
                pdf.set_x(indent)
                pdf.cell(10, 7, sanitize_text(f"{num_match.group(1)}."))
                pdf.set_font("Helvetica", "", 11)
                pdf.set_x(indent + 10)
                pdf.multi_cell(pdf.w - indent - 10 - pdf.r_margin, 7, sanitize_text(num_match.group(2)))
                pdf.ln(2)
                continue
            # Bullet items
            if line.startswith("- ") or line.startswith("* "):
                content = line[2:].strip()
                content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
                pdf.set_font("Helvetica", "", 11)
                indent = pdf.l_margin + 6
                pdf.set_x(indent)
                pdf.cell(6, 6, sanitize_text("-"))
                pdf.set_x(indent + 6)
                pdf.multi_cell(pdf.w - indent - 6 - pdf.r_margin, 6, sanitize_text(content))
                pdf.ln(1)
                continue
            # Default
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, sanitize_text(clean))
    else:
        pdf.set_font("Helvetica", "I", 11)
        pdf.cell(0, 8, sanitize_text("No clip breakdown found in script output."))

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(tmp.name)
    tmp.seek(0)
    with open(tmp.name, "rb") as f:
        return f.read()


def parse_clip_breakdown_to_docx(script_text: str) -> bytes:
    """Extract clip breakdown section and generate a DOCX."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title = "Untitled Script"
    title_match = re.search(r"(?:TITLE|Title)[:\s]*\n*(.+?)(?:\n|$)", script_text)
    if title_match:
        title = title_match.group(1).strip().strip("#").strip()

    clip_section = _extract_section(
        script_text, "CLIP BREAKDOWN",
        ["CASTING BREAKDOWN", "PROPS LIST", "LOCATION LIST"]
    )

    doc = Document()

    # Title
    heading = doc.add_heading("CLIP BREAKDOWN", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(title)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)

    doc.add_paragraph("")  # spacer

    if clip_section:
        for line in clip_section.split("\n"):
            line = line.strip()
            if not line:
                continue
            # Bold items
            bold_match = re.match(r"^\*\*(.+?)\*\*:?\s*(.*)", line)
            if bold_match:
                p = doc.add_paragraph()
                run = p.add_run(bold_match.group(1))
                run.bold = True
                run.font.size = Pt(11)
                if bold_match.group(2):
                    run2 = p.add_run(f": {bold_match.group(2)}")
                    run2.font.size = Pt(11)
                continue
            # Numbered items
            num_match = re.match(r"^(\d+)\.\s*(.*)", line)
            if num_match:
                p = doc.add_paragraph()
                run = p.add_run(f"{num_match.group(1)}. ")
                run.bold = True
                run.font.size = Pt(11)
                run2 = p.add_run(num_match.group(2))
                run2.font.size = Pt(11)
                continue
            # Bullet items
            if line.startswith("- ") or line.startswith("* "):
                content = line[2:].strip()
                content = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
                doc.add_paragraph(content, style="List Bullet")
                continue
            # Default
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            doc.add_paragraph(clean)
    else:
        doc.add_paragraph("No clip breakdown found in script output.", style="No Spacing")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    with open(tmp.name, "rb") as f:
        return f.read()


# ═══════════════════════════════════════════════════
# PROPS LIST extraction (for display)
# ═══════════════════════════════════════════════════

def extract_props_list(script_text: str) -> str:
    """Extract the props list section as plain text."""
    return _extract_section(script_text, "PROPS LIST", ["LOCATION LIST", "CASTING BREAKDOWN"])


def extract_casting_breakdown(script_text: str) -> str:
    """Extract the casting breakdown section as plain text."""
    return _extract_section(script_text, "CASTING BREAKDOWN", ["PROPS LIST", "LOCATION LIST"])


def extract_clip_breakdown(script_text: str) -> str:
    """Extract the clip breakdown section as plain text."""
    return _extract_section(
        script_text, "CLIP BREAKDOWN",
        ["CASTING BREAKDOWN", "PROPS LIST", "LOCATION LIST"]
    )


def extract_location_list(script_text: str) -> str:
    """Extract the location list section as plain text."""
    return _extract_section(script_text, "LOCATION LIST", ["CASTING BREAKDOWN", "PROPS LIST"])
